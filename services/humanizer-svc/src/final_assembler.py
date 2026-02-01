"""
Assembles the final verified DOCX document preserving all structure.
"""

from typing import List, Dict
from docx import Document
from docx.text.paragraph import Paragraph
from .document_parser import ParagraphInfo
from .structure_extractor import DocumentStructure


class FinalAssembler:
    """Assembles final verified DOCX document."""
    
    @staticmethod
    def assemble(
        original_doc: Document,
        structure: DocumentStructure,
        humanized_paragraphs: Dict[int, str]
    ) -> Document:
        """
        Assemble final document with humanized paragraphs.
        
        Args:
            original_doc: Original Document object
            structure: DocumentStructure with preserved structure
            humanized_paragraphs: Dict mapping paragraph index to humanized text
            
        Returns:
            New Document with humanized content
        """
        # Create new document
        new_doc = Document()
        
        # Copy document properties
        if original_doc.core_properties:
            new_doc.core_properties.title = original_doc.core_properties.title or ""
            new_doc.core_properties.author = original_doc.core_properties.author or ""
        
        # Create mapping from paragraph index to humanized text
        humanized_map = humanized_paragraphs
        
        # Create set of humanized paragraph indices for quick lookup
        humanized_indices = set(humanized_map.keys())
        
        # Create set of all paragraph indices that should be humanized
        para_indices_to_humanize = {p_info.index for p_info in structure.paragraphs}
        
        # Add all paragraphs in order
        for idx, para in enumerate(original_doc.paragraphs):
            # Check if this paragraph index should be humanized
            if idx in para_indices_to_humanize and idx in humanized_indices:
                # Use humanized text
                new_para = new_doc.add_paragraph()
                new_para.text = humanized_map[idx]
                # Copy formatting
                if para.style:
                    new_para.style = para.style
            else:
                # Copy original paragraph (heading, reference, etc.)
                new_para = new_doc.add_paragraph()
                new_para.text = para.text
                if para.style:
                    new_para.style = para.style
        
        # Copy tables
        for table in structure.tables:
            # Create new table with same dimensions
            new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            new_table.style = table.style
            
            # Copy table content
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_table.rows[i].cells[j].text = cell.text
        
        return new_doc

