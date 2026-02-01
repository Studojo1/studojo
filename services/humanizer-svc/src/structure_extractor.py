"""
Extract and preserve document structure: headings, tables, figures, references.
"""

from typing import List, Dict, Any, Set
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from .document_parser import ParagraphInfo


class DocumentStructure:
    """Represents the structure of a document."""
    
    def __init__(self):
        self.headings: List[ParagraphInfo] = []
        self.paragraphs: List[ParagraphInfo] = []
        self.tables: List[Table] = []
        self.table_positions: Dict[int, int] = {}  # paragraph index -> table index
        self.references: Set[str] = set()  # Reference patterns found
        self.figure_captions: List[str] = []


def is_heading(paragraph: Paragraph) -> bool:
    """
    Check if a paragraph is a heading.
    
    Args:
        paragraph: Paragraph object
        
    Returns:
        True if paragraph is a heading
    """
    if not paragraph.style:
        return False
    style_name = paragraph.style.name
    return style_name.startswith('Heading') or style_name.startswith('Title')


def is_reference(paragraph: Paragraph) -> bool:
    """
    Check if a paragraph contains a reference (citation).
    
    Args:
        paragraph: Paragraph object
        
    Returns:
        True if paragraph appears to be a reference
    """
    text = paragraph.text.strip()
    
    # Common reference patterns
    reference_patterns = [
        r'^\d+\.',  # Numbered reference: "1. Author..."
        r'^\[',     # Bracketed reference: "[1] Author..."
        r'^\(',     # Parenthesized: "(Author, Year)"
        r'^Author', # Starts with Author
        r'^References?', # References section header
        r'^Bibliography', # Bibliography header
    ]
    
    import re
    for pattern in reference_patterns:
        if re.match(pattern, text, re.IGNORECASE):
            return True
    
    return False


def is_figure_caption(paragraph: Paragraph) -> bool:
    """
    Check if a paragraph is a figure caption.
    
    Args:
        paragraph: Paragraph object
        
    Returns:
        True if paragraph appears to be a figure caption
    """
    text = paragraph.text.strip()
    
    # Common caption patterns
    caption_patterns = [
        r'^Figure \d+',
        r'^Fig\. \d+',
        r'^Table \d+',
        r'^Chart \d+',
    ]
    
    import re
    for pattern in caption_patterns:
        if re.match(pattern, text, re.IGNORECASE):
            return True
    
    return False


def extract_structure(document: Document, paragraphs: List[ParagraphInfo]) -> DocumentStructure:
    """
    Extract document structure, separating headings, tables, figures, and references.
    
    Args:
        document: Document object
        paragraphs: List of ParagraphInfo objects
        
    Returns:
        DocumentStructure object
    """
    structure = DocumentStructure()
    
    # Extract headings and regular paragraphs
    for para_info in paragraphs:
        para = para_info.paragraph
        
        if is_heading(para):
            structure.headings.append(para_info)
            para_info.is_heading = True
        elif is_reference(para):
            structure.references.add(para_info.text)
        elif is_figure_caption(para):
            structure.figure_captions.append(para_info.text)
        else:
            # Regular paragraph (candidate for humanization)
            structure.paragraphs.append(para_info)
    
    # Extract tables
    for idx, table in enumerate(document.tables):
        structure.tables.append(table)
        # Note: python-docx doesn't provide direct paragraph-to-table mapping
        # We'll preserve tables separately
    
    return structure


def get_paragraphs_to_humanize(structure: DocumentStructure) -> List[ParagraphInfo]:
    """
    Get list of paragraphs that should be humanized.
    Excludes headings, tables, figures, and references.
    
    Args:
        structure: DocumentStructure object
        
    Returns:
        List of ParagraphInfo objects to humanize
    """
    return structure.paragraphs

